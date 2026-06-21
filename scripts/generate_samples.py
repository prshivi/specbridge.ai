"""Generate deterministic binary parser samples for local development."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill

SAMPLES_DIR = Path(__file__).resolve().parents[1] / "samples"


def build_pdf() -> bytes:
    """Build a minimal one-page, text-based PDF without extra dependencies."""
    stream = (
        b"BT\n"
        b"/F1 18 Tf\n"
        b"72 730 Td\n"
        b"(SpecBridge PDF Sample) Tj\n"
        b"0 -32 Td\n"
        b"/F1 12 Tf\n"
        b"(Upload and parse business requirements.) Tj\n"
        b"ET\n"
    )
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        (
            b"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] "
            b"/Resources << /Font << /F1 5 0 R >> >> /Contents 4 0 R >>"
        ),
        b"<< /Length " + str(len(stream)).encode() + b" >>\nstream\n" + stream + b"endstream",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
    ]

    content = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for index, obj in enumerate(objects, start=1):
        offsets.append(len(content))
        content.extend(f"{index} 0 obj\n".encode())
        content.extend(obj)
        content.extend(b"\nendobj\n")

    xref_offset = len(content)
    content.extend(f"xref\n0 {len(objects) + 1}\n".encode())
    content.extend(b"0000000000 65535 f \n")
    for offset in offsets[1:]:
        content.extend(f"{offset:010d} 00000 n \n".encode())
    content.extend(
        (
            f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
            f"startxref\n{xref_offset}\n%%EOF\n"
        ).encode()
    )
    return bytes(content)


def create_docx(path: Path) -> None:
    """Create a compact business-style DOCX parser fixture."""
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)

    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(8)
    title_run = title.add_run("SpecBridge DOCX Sample")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(46, 116, 181)

    document.add_paragraph("Engineering-ready requirements parser fixture.")
    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.rows[0].cells[0].text = "Requirement ID"
    table.rows[0].cells[1].text = "Title"
    row = table.add_row().cells
    row[0].text = "REQ-001"
    row[1].text = "Upload documents"
    document.save(path)


def create_xlsx(path: Path) -> None:
    """Create a small workbook parser fixture."""
    workbook = Workbook()
    worksheet = workbook.active
    worksheet.title = "Requirements"
    worksheet.append(["Requirement ID", "Title", "Priority"])
    worksheet.append(["REQ-001", "Upload documents", "High"])
    worksheet.append(["REQ-002", "Parse text", "High"])
    worksheet.append(["REQ-003", "Store originals", "Medium"])

    header_fill = PatternFill("solid", fgColor="0F766E")
    for cell in worksheet[1]:
        cell.font = Font(color="FFFFFF", bold=True)
        cell.fill = header_fill
    worksheet.freeze_panes = "A2"
    worksheet.column_dimensions["A"].width = 18
    worksheet.column_dimensions["B"].width = 28
    worksheet.column_dimensions["C"].width = 14
    workbook.save(path)


def main() -> None:
    SAMPLES_DIR.mkdir(parents=True, exist_ok=True)
    (SAMPLES_DIR / "sample.pdf").write_bytes(build_pdf())
    create_docx(SAMPLES_DIR / "sample.docx")
    create_xlsx(SAMPLES_DIR / "sample.xlsx")


if __name__ == "__main__":
    main()
